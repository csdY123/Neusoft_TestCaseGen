"""
RAG utility module for knowledge retrieval
Supports both local HuggingFace embedding and Ollama API embedding
Includes LLM-based document chunking
"""

import os

# Force offline mode to avoid network issues
os.environ["HF_HUB_OFFLINE"] = "1"
os.environ["TRANSFORMERS_OFFLINE"] = "1"
import json
import logging
import torch
import numpy as np
import jieba
import re
import requests
from typing import List, Tuple, Dict, Any, Optional, Callable
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from rank_bm25 import BM25Okapi
from langdetect import detect
from tqdm import tqdm
from docx import Document as DocxDocument
from openai import OpenAI

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# LLM config for chunking (uses local vLLM by default)
LLM_CONFIG = {
    "api_key": "EMPTY",
    "base_url": "http://localhost:12349/v1",
    "model": "Qwen3-8B"
}

# Embedding config - supports both local and API modes
EMBEDDING_CONFIG = {
    "mode": "local",  # "local" or "api"
    "api_base_url": "http://localhost:11434",  # Ollama API base URL
    "api_model": "bge-large",  # Ollama embedding model name
    "local_model_path": "/media/a100/c5e1bf65-7974-432f-8aed-7a1345241efe/chensenda/codes/models/bge-large-zh-v1.5"
}

# Global instances
_embeddings = None
_vectorstore = None
_reranker = None


class LocalEmbeddings:
    """Local BGE embedding model wrapper"""
    
    # Default to local cached path
    DEFAULT_MODEL_PATH = os.path.expanduser(
        "/media/a100/c5e1bf65-7974-432f-8aed-7a1345241efe/chensenda/codes/models/bge-large-zh-v1.5"
    )
    
    def __init__(self, model_name: str = None, device: str = "cuda"):
        model_path = model_name or self.DEFAULT_MODEL_PATH
        self.model = HuggingFaceEmbeddings(
            model_name=model_path,
            model_kwargs={'device': device},
            encode_kwargs={'normalize_embeddings': True}
        )
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self.model.embed_documents(texts)
    
    def embed_query(self, text: str) -> List[float]:
        return self.model.embed_query(text)


class OllamaEmbeddings:
    """Ollama API embedding model wrapper"""
    
    def __init__(self, base_url: str = "http://localhost:11434", model: str = "bge-large"):
        self.base_url = base_url.rstrip("/")
        self.model = model
        self._dimension = None
        logger.info(f"Initialized OllamaEmbeddings with model: {model} at {base_url}")
    
    def _get_embedding(self, text: str) -> List[float]:
        """Get embedding for a single text via Ollama API"""
        url = f"{self.base_url}/api/embed"
        payload = {
            "model": self.model,
            "input": text
        }
        try:
            response = requests.post(url, json=payload, timeout=60)
            response.raise_for_status()
            result = response.json()
            embeddings = result.get("embeddings", [])
            if embeddings and len(embeddings) > 0:
                return embeddings[0]
            raise ValueError("No embeddings returned from Ollama API")
        except requests.exceptions.RequestException as e:
            logger.error(f"Ollama API request failed: {e}")
            raise
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents"""
        embeddings = []
        for text in texts:
            embedding = self._get_embedding(text)
            embeddings.append(embedding)
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a query text"""
        return self._get_embedding(text)


class EmbeddingsWrapper(Embeddings):
    """Wrapper class that provides unified interface for different embedding backends.
    Compatible with LangChain FAISS vectorstore.
    Inherits from LangChain Embeddings base class for proper compatibility.
    """
    
    def __init__(self, embeddings_impl):
        self._impl = embeddings_impl
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents."""
        return self._impl.embed_documents(texts)
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a single query text."""
        return self._impl.embed_query(text)


class Qwen3Reranker:
    """Qwen3 Reranker for document reranking"""
    
    def __init__(self, model_path: str, device: str = None):
        from transformers import AutoTokenizer
        from modelscope import AutoModelForCausalLM as MSAutoModelForCausalLM
        
        logger.info(f"Loading Qwen3 Reranker from: {model_path}")
        self.tokenizer = AutoTokenizer.from_pretrained(
            model_path, padding_side='left', trust_remote_code=True
        )
        
        self.device = device or ("cuda" if torch.cuda.is_available() else "cpu")
        
        try:
            self.model = MSAutoModelForCausalLM.from_pretrained(
                model_path,
                torch_dtype=torch.float16,
                attn_implementation="flash_attention_2",
                trust_remote_code=True
            ).to(self.device).eval()
            logger.info("Reranker loaded with flash attention")
        except Exception as e:
            logger.warning(f"Flash attention unavailable: {e}")
            self.model = MSAutoModelForCausalLM.from_pretrained(
                model_path, trust_remote_code=True
            ).to(self.device).eval()
        
        self.token_false_id = self.tokenizer.convert_tokens_to_ids("no")
        self.token_true_id = self.tokenizer.convert_tokens_to_ids("yes")
        self.max_length = 8192
        
        prefix = "<|im_start|>system\nJudge whether the Document meets the requirements based on the Query and the Instruct provided. Note that the answer can only be \"yes\" or \"no\".<|im_end|>\n<|im_start|>user\n"
        suffix = "<|im_end|>\n<|im_start|>assistant\n<think>\n\n</think>\n\n"
        self.prefix_tokens = self.tokenizer.encode(prefix, add_special_tokens=False)
        self.suffix_tokens = self.tokenizer.encode(suffix, add_special_tokens=False)
    
    def format_instruction(self, query: str, doc: str) -> str:
        instruction = 'Given a web search query, retrieve relevant passages that answer the query'
        return f"<Instruct>: {instruction}\n<Query>: {query}\n<Document>: {doc}"
    
    def process_inputs(self, pairs: List[str]):
        inputs = self.tokenizer(
            pairs, padding=False, truncation='longest_first',
            return_attention_mask=False,
            max_length=self.max_length - len(self.prefix_tokens) - len(self.suffix_tokens)
        )
        for i, ele in enumerate(inputs['input_ids']):
            inputs['input_ids'][i] = self.prefix_tokens + ele + self.suffix_tokens
        inputs = self.tokenizer.pad(inputs, padding=True, return_tensors="pt", max_length=self.max_length)
        for key in inputs:
            inputs[key] = inputs[key].to(self.device)
        return inputs
    
    @torch.no_grad()
    def compute_scores(self, query: str, docs: List[str], batch_size: int = 8) -> List[float]:
        pairs = [self.format_instruction(query, doc) for doc in docs]
        scores = []
        
        for i in range(0, len(pairs), batch_size):
            batch_pairs = pairs[i:i + batch_size]
            inputs = self.process_inputs(batch_pairs)
            batch_scores = self.model(**inputs).logits[:, -1, :]
            true_vector = batch_scores[:, self.token_true_id]
            false_vector = batch_scores[:, self.token_false_id]
            batch_scores = torch.stack([false_vector, true_vector], dim=1)
            batch_scores = torch.nn.functional.log_softmax(batch_scores, dim=1)
            batch_scores = batch_scores[:, 1].exp().tolist()
            scores.extend(batch_scores)
        
        return scores


def configure_embeddings(
    mode: str = None,
    api_base_url: str = None,
    api_model: str = None,
    local_model_path: str = None
):
    """Configure embedding model settings.
    
    Args:
        mode: "local" for HuggingFace local model, "api" for Ollama API
        api_base_url: Ollama API base URL (for api mode)
        api_model: Ollama embedding model name (for api mode)
        local_model_path: Path to local HuggingFace model (for local mode)
    """
    global _embeddings
    
    if mode is not None:
        EMBEDDING_CONFIG["mode"] = mode
    if api_base_url is not None:
        EMBEDDING_CONFIG["api_base_url"] = api_base_url
    if api_model is not None:
        EMBEDDING_CONFIG["api_model"] = api_model
    if local_model_path is not None:
        EMBEDDING_CONFIG["local_model_path"] = local_model_path
    
    # Reset embeddings instance to apply new config
    _embeddings = None
    logger.info(f"Embedding config updated: mode={EMBEDDING_CONFIG['mode']}, "
                f"api_model={EMBEDDING_CONFIG['api_model']}")


def get_embeddings() -> EmbeddingsWrapper:
    """Get or create embedding model instance based on current config"""
    global _embeddings
    if _embeddings is None:
        mode = EMBEDDING_CONFIG.get("mode", "local")
        
        if mode == "api":
            logger.info(f"Initializing Ollama API embedding model: {EMBEDDING_CONFIG['api_model']}")
            impl = OllamaEmbeddings(
                base_url=EMBEDDING_CONFIG["api_base_url"],
                model=EMBEDDING_CONFIG["api_model"]
            )
            _embeddings = EmbeddingsWrapper(impl)
            logger.info("Ollama embedding model initialized")
        else:
            logger.info("Initializing local HuggingFace embedding model...")
            impl = LocalEmbeddings(model_name=EMBEDDING_CONFIG.get("local_model_path"))
            _embeddings = EmbeddingsWrapper(impl)
            logger.info("Local embedding model initialized")
    
    return _embeddings


def get_embedding_config() -> Dict[str, Any]:
    """Get current embedding configuration"""
    return EMBEDDING_CONFIG.copy()


def load_vectorstore(index_directory: str = "faiss_index") -> Optional[FAISS]:
    """Load FAISS vectorstore"""
    global _vectorstore
    
    index_path = os.path.join(index_directory, "index.faiss")
    if not os.path.exists(index_path):
        logger.warning(f"FAISS index not found at {index_directory}")
        return None
    
    if _vectorstore is None:
        logger.info(f"Loading FAISS index from {index_directory}")
        embeddings = get_embeddings()
        # Use the wrapper's model attribute which points to itself
        _vectorstore = FAISS.load_local(
            index_directory, embeddings, allow_dangerous_deserialization=True
        )
        logger.info("FAISS index loaded")
    
    return _vectorstore


def get_reranker(model_path: str = None) -> Optional[Qwen3Reranker]:
    """Get or create reranker instance"""
    global _reranker
    
    if model_path is None:
        model_path = "/media/a100/c5e1bf65-7974-432f-8aed-7a1345241efe/zts/.cache/modelscope/hub/models/Qwen/Qwen3-Reranker-0.6B"
    
    if not os.path.exists(model_path):
        logger.warning(f"Reranker model not found at {model_path}")
        return None
    
    if _reranker is None:
        _reranker = Qwen3Reranker(model_path)
    
    return _reranker


def hybrid_search(
    query: str,
    vectorstore: FAISS,
    top_k: int = 10,
    alpha: float = 0.8,
    recall_factor: int = 3,
    reranker: Qwen3Reranker = None
) -> List[Tuple[Document, float]]:
    """Hybrid search combining vector search and BM25, with optional reranking"""
    
    initial_top_k = top_k * recall_factor
    
    # Vector search
    vector_results = vectorstore.similarity_search_with_score(query, k=initial_top_k)
    
    # Get all docs for BM25
    all_docs = list(vectorstore.docstore._dict.values())
    docs_content = [doc.page_content for doc in all_docs]
    
    # Language detection for tokenization
    try:
        lang = detect(query)
    except:
        lang = 'zh'
    
    if lang == 'en':
        tokenized_corpus = [re.findall(r'\w+', doc.lower()) for doc in docs_content]
        tokenized_query = re.findall(r'\w+', query.lower())
    else:
        tokenized_corpus = [list(jieba.cut(doc)) for doc in docs_content]
        tokenized_query = list(jieba.cut(query))
    
    # BM25 search
    bm25 = BM25Okapi(tokenized_corpus)
    bm25_scores = bm25.get_scores(tokenized_query)
    keyword_indices = bm25_scores.argsort()[-initial_top_k:][::-1]
    keyword_results = [(all_docs[i], bm25_scores[i]) for i in keyword_indices]
    
    # Normalize scores
    vector_max = max([score for _, score in vector_results]) if vector_results else 1.0
    keyword_max = max([score for _, score in keyword_results]) if keyword_results else 1.0
    if vector_max == 0: vector_max = 1.0
    if keyword_max == 0: keyword_max = 1.0
    
    # Merge results
    vector_scores = {hash(doc.page_content): (1 - score / vector_max) for doc, score in vector_results}
    keyword_scores = {hash(doc.page_content): score / keyword_max for doc, score in keyword_results}
    
    all_docs_map = {}
    for doc, _ in vector_results:
        all_docs_map[hash(doc.page_content)] = doc
    for doc, _ in keyword_results:
        all_docs_map[hash(doc.page_content)] = doc
    
    # Calculate hybrid scores
    hybrid_results = []
    for doc_hash, doc in all_docs_map.items():
        vector_score = vector_scores.get(doc_hash, 0)
        keyword_score = keyword_scores.get(doc_hash, 0)
        hybrid_score = alpha * vector_score + (1 - alpha) * keyword_score
        hybrid_results.append((doc, hybrid_score))
    
    hybrid_results.sort(key=lambda x: x[1], reverse=True)
    initial_results = hybrid_results[:initial_top_k]
    
    # Rerank if available
    if reranker:
        logger.info("Reranking with Qwen3 Reranker...")
        docs_content = [doc.page_content for doc, _ in initial_results]
        rerank_scores = reranker.compute_scores(query, docs_content)
        reranked_results = [(initial_results[i][0], float(rerank_scores[i])) for i in range(len(initial_results))]
        reranked_results.sort(key=lambda x: x[1], reverse=True)
        return reranked_results[:top_k]
    
    return initial_results[:top_k]


def retrieve_knowledge(
    query: str,
    top_k: int = 7,
    use_reranker: bool = False,
    index_directory: str = "faiss_index"
) -> List[str]:
    """Retrieve relevant knowledge fragments for a query"""
    
    vectorstore = load_vectorstore(index_directory)
    if vectorstore is None:
        return []
    
    reranker = get_reranker() if use_reranker else None
    
    results = hybrid_search(query, vectorstore, top_k=top_k, reranker=reranker)
    return [doc.page_content for doc, _ in results]


def format_retrieved_content(fragments: List[str], query: str = None) -> str:
    """Format retrieved fragments as PRD content"""
    if not fragments:
        return ""
    
    content = "Retrieved Knowledge Fragments:\n\n"
    for i, fragment in enumerate(fragments, 1):
        content += f"【Fragment {i}】\n{fragment}\n\n"
    
    return content


# ========== Document Chunking Functions ==========

def infer_heading_level_by_font(docx_path: str) -> List[Dict]:
    """Infer heading levels by font size and indent"""
    doc = DocxDocument(docx_path)
    font_sizes = set()
    indent_levels = set()
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.add(run.font.size.pt)
        if para.paragraph_format.left_indent:
            indent_levels.add(para.paragraph_format.left_indent)
    
    sorted_font_sizes = sorted(font_sizes, reverse=True)
    sorted_indents = sorted(indent_levels, reverse=True)
    
    structured_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        current_font_size = None
        for run in para.runs:
            if run.font.size:
                current_font_size = run.font.size.pt
                break
        
        current_indent = para.paragraph_format.left_indent
        
        font_level = 0
        if current_font_size and sorted_font_sizes:
            font_level = sorted_font_sizes.index(current_font_size) + 1
        
        indent_level = 0
        if current_indent and sorted_indents:
            indent_level = sorted_indents.index(current_indent) + 1
        
        final_level = max(font_level, indent_level)
        structured_text.append({"text": text, "level": final_level})
    
    return structured_text


def docx_to_markdown(docx_path: str) -> Tuple[List[str], int]:
    """Convert docx to markdown-like structure"""
    structured_text = infer_heading_level_by_font(docx_path)
    if not structured_text:
        return [], 1
    
    max_level = max(item["level"] for item in structured_text)
    markdown_lines = []
    
    for para in structured_text:
        level = para['level']
        text = para['text'].strip()
        if level > 0 and level < max_level:
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            markdown_lines.append(text)
    
    return markdown_lines, max_level


def docx_table_to_markdown(docx_path: str) -> List[List[str]]:
    """Extract tables from docx as markdown"""
    doc = DocxDocument(docx_path)
    tables = []
    
    for table in doc.tables:
        tb = []
        tb.append("\n| " + " | ".join(["Header"] * len(table.columns)) + " |")
        tb.append("| " + " | ".join(["---"] * len(table.columns)) + " |")
        for row in table.rows:
            tb.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")
        tables.append(tb)
    
    return tables


def llm_segment(client: OpenAI, content: str, model: str) -> Optional[str]:
    """Use LLM to segment document content"""
    system_prompt = """You are a document segmentation assistant. Segment the input into thematic chunks.

Input format: Text fragments with [EX N] prefix.
Output format: JSONL only, one JSON per line, no other text.

Each JSON must have:
- segment_id: index from 0
- start_exchange_number: first fragment number
- end_exchange_number: last fragment number  
- num_exchange: count of fragments

Example output:
{"segment_id": 0, "start_exchange_number": 0, "end_exchange_number": 5, "num_exchange": 6}
{"segment_id": 1, "start_exchange_number": 6, "end_exchange_number": 8, "num_exchange": 3}

Rules:
1. Cover ALL fragments, no gaps allowed
2. Each segment should be thematically coherent
3. Output ONLY JSON lines, nothing else"""

    # Add /no_think for Qwen3 to disable thinking mode
    user_content = content + "\n\n/no_think"
    
    try:
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            temperature=0.7,
            max_tokens=4096,
        )
        return completion.choices[0].message.content.strip() if completion.choices[0].message.content else None
    except Exception as e:
        logger.error(f"LLM segment error: {e}")
        return None


def document_level_segment(content_list: List[str], max_level: int, ex_flag: bool = True) -> List[List[str]]:
    """Coarse-grained document segmentation by heading levels"""
    full_texts = []
    para = []
    prefix = [f"{'#' * l}" for l in range(1, max_level)]
    
    for i, c in enumerate(content_list):
        item = f"[EX {i}] {c}" if ex_flag else c
        c = c.strip()
        
        para_flag = any(c.startswith(pre) for pre in prefix)
        if para_flag:
            full_texts.append(para)
            para = [item]
        else:
            para.append(item)
    
    full_texts.append(para)
    return full_texts


def find_uncovered_indices(covered_ranges: List[Dict], total_end: int) -> List[Tuple[int, int]]:
    """Find indices not covered by any segment"""
    merged_ranges = []
    for current in sorted(covered_ranges, key=lambda x: x["start_exchange_number"]):
        if not merged_ranges:
            merged_ranges.append(current)
        else:
            last = merged_ranges[-1]
            if int(current["start_exchange_number"]) <= int(last["end_exchange_number"]) + 1:
                last["end_exchange_number"] = max(int(last["end_exchange_number"]), int(current["end_exchange_number"]))
            else:
                merged_ranges.append(current)
    
    uncovered = []
    prev_end = -1
    
    for interval in merged_ranges:
        start_no = int(interval["start_exchange_number"])
        end_no = int(interval["end_exchange_number"])
        if start_no > prev_end + 1:
            uncovered.append((prev_end + 1, start_no - 1))
        prev_end = max(prev_end, end_no)
    
    if prev_end < total_end:
        uncovered.append((prev_end + 1, total_end))
    
    return uncovered


def deduplicate_chunks(chunks: List[Dict], min_len: int = 10) -> List[Dict]:
    """Deduplicate and sort chunks"""
    seen = set()
    deduplicated = []
    
    for chunk in chunks:
        if len(chunk.get('text', '')) <= min_len:
            continue
        key = (int(chunk["start_exchange_number"]), int(chunk["end_exchange_number"]))
        if key not in seen:
            seen.add(key)
            deduplicated.append(chunk)
    
    sorted_chunks = sorted(deduplicated, key=lambda x: (int(x["start_exchange_number"]), int(x['end_exchange_number'])))
    for i, chunk in enumerate(sorted_chunks):
        chunk["segment_id"] = i
    
    return sorted_chunks


def chunk_docx_fast(docx_path: str, progress_callback=None, max_chunk_size: int = 800) -> List[Dict[str, Any]]:
    """
    Fast document chunking using rule-based segmentation (no LLM calls).
    Splits by heading structure and paragraph boundaries.
    """
    if progress_callback:
        progress_callback("Parsing document structure...")
    
    content_list, max_level = docx_to_markdown(docx_path)
    table_list = docx_table_to_markdown(docx_path)
    
    if not content_list and not table_list:
        return []
    
    if progress_callback:
        progress_callback(f"Found {len(content_list)} paragraphs, {len(table_list)} tables")
    
    # Segment by heading structure
    chunks = []
    current_chunk = []
    current_size = 0
    
    for line in content_list:
        line_size = len(line)
        
        # Check if this is a heading (starts with #)
        is_heading = line.strip().startswith('#')
        
        # Start new chunk if: heading found OR current chunk too large
        if is_heading and current_chunk:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]
            current_size = line_size
        elif current_size + line_size > max_chunk_size and current_chunk:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]
            current_size = line_size
        else:
            current_chunk.append(line)
            current_size += line_size
    
    # Add remaining content
    if current_chunk:
        chunks.append("\n".join(current_chunk))
    
    # Add tables as separate chunks
    for tb in table_list:
        tb_txt = "\n".join(tb)
        chunks.append(tb_txt)
    
    if progress_callback:
        progress_callback(f"Created {len(chunks)} chunks (fast mode)")
    
    # Convert to standard format
    result = []
    for i, chunk_text in enumerate(chunks):
        if len(chunk_text.strip()) > 10:  # Filter empty chunks
            result.append({
                "id": str(i),
                "content": chunk_text,
                "available": True
            })
    
    if progress_callback:
        progress_callback(f"Done! {len(result)} valid chunks")
    
    return result


def chunk_docx_with_llm(docx_path: str, progress_callback=None, use_llm: bool = False) -> List[Dict[str, Any]]:
    """
    Chunk a docx file. 
    - use_llm=False (default): Fast rule-based chunking
    - use_llm=True: LLM-based semantic segmentation (slower, batched)
    """
    # Default to fast mode
    if not use_llm:
        return chunk_docx_fast(docx_path, progress_callback)
    
    # LLM mode - batch processing to reduce API calls
    if progress_callback:
        progress_callback("Parsing document structure...")
    
    content_list, max_level = docx_to_markdown(docx_path)
    table_list = docx_table_to_markdown(docx_path)
    
    if not content_list and not table_list:
        return []
    
    if progress_callback:
        progress_callback(f"Found {len(content_list)} paragraphs, {len(table_list)} tables")
    
    # Create numbered content for LLM
    numbered_content = []
    for i, line in enumerate(content_list):
        numbered_content.append(f"[EX {i}] {line}")
    
    # Batch into ~6000 char chunks to fit context window
    batches = []
    current_batch = []
    current_size = 0
    batch_limit = 6000
    
    for line in numbered_content:
        if current_size + len(line) > batch_limit and current_batch:
            batches.append("\n\n".join(current_batch))
            current_batch = [line]
            current_size = len(line)
        else:
            current_batch.append(line)
            current_size += len(line)
    if current_batch:
        batches.append("\n\n".join(current_batch))
    
    if progress_callback:
        progress_callback(f"Processing {len(batches)} batches with LLM...")
    
    # Call LLM for each batch (much fewer calls now)
    client = OpenAI(api_key=LLM_CONFIG["api_key"], base_url=LLM_CONFIG["base_url"])
    
    seg_ids = []
    for i, batch in enumerate(batches):
        if progress_callback:
            progress_callback(f"LLM batch {i+1}/{len(batches)}...")
        seg_json = llm_segment(client, batch, LLM_CONFIG["model"])
        if seg_json:
            seg_ids += seg_json.strip().split("\n")
    
    # Parse LLM results - handle various formats
    llm_seg_results = []
    all_text = "\n".join(seg_ids)
    
    # Remove thinking tags if present (Qwen3 format)
    all_text = re.sub(r'<think>.*?</think>', '', all_text, flags=re.DOTALL)
    all_text = re.sub(r'<\|.*?\|>', '', all_text)  # Remove special tokens
    
    # Find all JSON objects
    json_pattern = r'\{[^{}]*"segment_id"[^{}]*\}'
    matches = re.findall(json_pattern, all_text)
    
    for match in matches:
        try:
            seg = json.loads(match)
            if 'start_exchange_number' in seg and 'end_exchange_number' in seg:
                llm_seg_results.append(seg)
        except json.JSONDecodeError:
            continue
    
    if progress_callback:
        progress_callback(f"LLM returned {len(llm_seg_results)} valid segments")
    
    # Build chunks from segments
    chunks = []
    for seg in llm_seg_results:
        start = int(seg.get('start_exchange_number', 0))
        end = int(seg.get('end_exchange_number', start))
        if start < len(content_list) and end < len(content_list) and start >= 0 and end >= 0:
            merge_text = "\n".join(content_list[start:end + 1])
            chunks.append({"text": merge_text, "start": start, "end": end})
    
    # Handle uncovered ranges - ALWAYS check even if chunks is empty
    if content_list:
        if chunks:
            covered = set()
            for c in chunks:
                for i in range(c["start"], c["end"] + 1):
                    covered.add(i)
            for i in range(len(content_list)):
                if i not in covered:
                    chunks.append({"text": content_list[i], "start": i, "end": i})
        else:
            # LLM returned no valid results, fallback to adding all content
            if progress_callback:
                progress_callback("LLM returned no valid segments, using fallback...")
            # Use fast mode as fallback
            return chunk_docx_fast(docx_path, progress_callback)
    
    # Add tables
    for tb in table_list:
        chunks.append({"text": "\n".join(tb), "start": -1, "end": -1})
    
    # Convert to standard format
    result = []
    for i, chunk in enumerate(chunks):
        if len(chunk["text"].strip()) > 10:
            result.append({
                "id": str(i),
                "content": chunk["text"],
                "available": True
            })
    
    if progress_callback:
        progress_callback(f"Done! {len(result)} chunks")
    
    return result


# ========== Index Building Functions ==========

def load_chunks_from_json(json_file_path: str) -> List[Dict[str, Any]]:
    """Load chunks from JSON file"""
    logger.info(f"Loading chunks from {json_file_path}")
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        logger.error(f"Failed to load JSON: {e}")
        return []
    
    chunks = []
    for document in data.get("documents", []):
        if "chunks" in document:
            for chunk in document["chunks"]:
                if chunk.get("available", True) and chunk.get("content"):
                    chunk["id"] = chunk.get("id", "")
                    chunks.append(chunk)
    
    logger.info(f"Loaded {len(chunks)} chunks")
    return chunks


def preprocess_chunks(chunks: List[Dict[str, Any]], min_length: int = 10) -> List[Dict[str, Any]]:
    """Preprocess chunks, filter empty/short content"""
    processed = []
    for chunk in chunks:
        if not chunk.get("content") or len(chunk["content"]) < min_length:
            continue
        processed.append(chunk)
    return processed


def convert_to_documents(chunks: List[Dict[str, Any]]) -> List[Document]:
    """Convert chunks to LangChain Documents"""
    documents = []
    for chunk in chunks:
        page_content = chunk["content"]
        metadata = {k: v for k, v in chunk.items() if k != "content"}
        documents.append(Document(page_content=page_content, metadata=metadata))
    return documents


def build_faiss_index(
    json_file_path: str,
    index_directory: str = "faiss_index",
    batch_size: int = 32
) -> bool:
    """Build FAISS index from JSON file"""
    
    os.makedirs(index_directory, exist_ok=True)
    
    # Load and preprocess
    raw_chunks = load_chunks_from_json(json_file_path)
    if not raw_chunks:
        logger.error("No chunks found")
        return False
    
    chunks = preprocess_chunks(raw_chunks)
    logger.info(f"Preprocessed: {len(raw_chunks)} -> {len(chunks)} chunks")
    
    documents = convert_to_documents(chunks)
    embeddings = get_embeddings()
    
    # Build index in batches
    logger.info("Building FAISS index...")
    
    first_batch = documents[:min(batch_size, len(documents))]
    vectorstore = FAISS.from_documents(first_batch, embeddings)
    remaining = documents[min(batch_size, len(documents)):]
    
    total_batches = (len(remaining) + batch_size - 1) // batch_size
    for i in tqdm(range(total_batches), desc="Building index"):
        start_idx = i * batch_size
        end_idx = min((i + 1) * batch_size, len(remaining))
        batch = remaining[start_idx:end_idx]
        vectorstore.add_documents(batch)
    
    vectorstore.save_local(index_directory)
    logger.info(f"Index saved to {index_directory}")
    
    # Update global instance
    global _vectorstore
    _vectorstore = vectorstore
    
    return True


def check_index_exists(index_directory: str = "faiss_index") -> bool:
    """Check if FAISS index exists"""
    return os.path.exists(os.path.join(index_directory, "index.faiss"))


def build_index_from_docx(
    docx_path: str,
    index_directory: str = "faiss_index",
    progress_callback=None
) -> bool:
    """Build index using default fast mode"""
    return build_index_from_docx_with_mode(docx_path, index_directory, use_llm=False, progress_callback=progress_callback)


def build_index_from_docx_with_mode(
    docx_path: str,
    index_directory: str = "faiss_index",
    use_llm: bool = False,
    progress_callback=None
) -> bool:
    """
    Build or update FAISS index from a single docx file.
    use_llm=False: Fast rule-based chunking
    use_llm=True: LLM semantic chunking
    """
    if not os.path.exists(docx_path):
        logger.error(f"File not found: {docx_path}")
        return False
    
    mode_name = "LLM" if use_llm else "Fast"
    if progress_callback:
        progress_callback(f"Chunking document ({mode_name} mode)...")
    
    # Chunk document
    chunks = chunk_docx_with_llm(docx_path, progress_callback, use_llm=use_llm)
    if not chunks:
        logger.error("No chunks generated")
        return False
    
    # Preprocess
    processed = preprocess_chunks(chunks)
    if not processed:
        logger.error("No valid chunks after preprocessing")
        return False
    
    if progress_callback:
        progress_callback(f"Building index with {len(processed)} chunks...")
    
    # Convert to documents
    documents = convert_to_documents(processed)
    embeddings = get_embeddings()
    
    os.makedirs(index_directory, exist_ok=True)
    
    # Check existing index
    global _vectorstore
    index_path = os.path.join(index_directory, "index.faiss")
    
    if os.path.exists(index_path):
        if progress_callback:
            progress_callback("Loading existing index and adding documents...")
        try:
            vectorstore = FAISS.load_local(
                index_directory, embeddings, allow_dangerous_deserialization=True
            )
            vectorstore.add_documents(documents)
        except Exception as e:
            logger.warning(f"Failed to load existing index: {e}, creating new one")
            vectorstore = FAISS.from_documents(documents, embeddings)
    else:
        if progress_callback:
            progress_callback("Creating new index...")
        vectorstore = FAISS.from_documents(documents, embeddings)
    
    vectorstore.save_local(index_directory)
    _vectorstore = vectorstore
    
    if progress_callback:
        progress_callback(f"✅ Index saved! Total chunks: {len(documents)}")
    
    return True


def get_index_stats(index_directory: str = "faiss_index") -> Dict[str, Any]:
    """Get statistics about the current index (without loading model)"""
    if not check_index_exists(index_directory):
        return {"exists": False, "num_documents": 0}
    
    # Just check file exists, don't load model at startup
    # Actual count will be shown after first use
    pkl_path = os.path.join(index_directory, "index.pkl")
    if os.path.exists(pkl_path):
        try:
            import pickle
            with open(pkl_path, "rb") as f:
                data = pickle.load(f)
                if hasattr(data, '__len__'):
                    return {"exists": True, "num_documents": len(data)}
        except:
            pass
    
    return {"exists": True, "num_documents": -1}  # -1 means unknown count


# ========== JSONL RAG for UI Automation Test Cases ==========

# Global cache for JSONL data
_jsonl_data_cache = {}
_jsonl_vectorstore = {}


def load_jsonl_data(jsonl_path: str) -> List[Dict[str, Any]]:
    """Load JSONL file and cache the data"""
    global _jsonl_data_cache
    
    if jsonl_path in _jsonl_data_cache:
        return _jsonl_data_cache[jsonl_path]
    
    if not os.path.exists(jsonl_path):
        logger.warning(f"JSONL file not found: {jsonl_path}")
        return []
    
    data = []
    try:
        with open(jsonl_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line:
                    try:
                        item = json.loads(line)
                        data.append(item)
                    except json.JSONDecodeError as e:
                        logger.warning(f"Failed to parse line: {e}")
                        continue
        _jsonl_data_cache[jsonl_path] = data
        logger.info(f"Loaded {len(data)} items from {jsonl_path}")
    except Exception as e:
        logger.error(f"Failed to load JSONL: {e}")
        return []
    
    return data


def build_jsonl_index(jsonl_path: str, index_name: str = "jsonl_index") -> bool:
    """Build FAISS index from JSONL file for RAG retrieval"""
    global _jsonl_vectorstore
    
    data = load_jsonl_data(jsonl_path)
    if not data:
        return False
    
    # Create documents from JSONL entries
    # Use prd_info fields as searchable content
    documents = []
    for i, item in enumerate(data):
        prd_info = item.get("prd_info", {})
        
        # Combine searchable fields
        search_text = " ".join([
            prd_info.get("test_case_name", ""),
            prd_info.get("test_point", ""),
            prd_info.get("feature", ""),
            prd_info.get("prd_document", "")
        ])
        
        if search_text.strip():
            metadata = {
                "index": i,
                "episode_id": item.get("episode_id", ""),
                "source": jsonl_path
            }
            documents.append(Document(page_content=search_text, metadata=metadata))
    
    if not documents:
        logger.error("No valid documents to index")
        return False
    
    embeddings = get_embeddings()
    vectorstore = FAISS.from_documents(documents, embeddings)
    _jsonl_vectorstore[index_name] = vectorstore
    
    logger.info(f"Built JSONL index with {len(documents)} documents")
    return True


def retrieve_jsonl_examples(
    query: str,
    jsonl_path: str,
    top_k: int = 3,
    index_name: str = "jsonl_index"
) -> List[Dict[str, Any]]:
    """
    Retrieve similar examples from JSONL file using RAG.
    Returns full JSONL entries (including steps) for few-shot learning.
    """
    global _jsonl_vectorstore
    
    # Build index if not exists
    if index_name not in _jsonl_vectorstore:
        if not build_jsonl_index(jsonl_path, index_name):
            return []
    
    vectorstore = _jsonl_vectorstore[index_name]
    data = load_jsonl_data(jsonl_path)
    
    if not data:
        return []
    
    # Search for similar documents
    try:
        results = vectorstore.similarity_search_with_score(query, k=top_k)
    except Exception as e:
        logger.error(f"Search failed: {e}")
        return []
    
    # Get full entries from original data
    examples = []
    seen_indices = set()
    
    for doc, score in results:
        idx = doc.metadata.get("index")
        if idx is not None and idx < len(data) and idx not in seen_indices:
            seen_indices.add(idx)
            examples.append(data[idx])
    
    logger.info(f"Retrieved {len(examples)} examples for query")
    return examples


def format_jsonl_examples_for_prompt(examples: List[Dict[str, Any]]) -> str:
    """Format retrieved JSONL examples as few-shot learning examples for prompt"""
    if not examples:
        return ""
    
    formatted = "\n## Reference Examples:\n"
    formatted += "Here are similar test cases for reference:\n\n"
    
    for i, example in enumerate(examples, 1):
        prd_info = example.get("prd_info", {})
        steps = example.get("steps", [])
        
        formatted += f"### Example {i}:\n"
        formatted += f"- **Test Case Name**: {prd_info.get('test_case_name', 'N/A')}\n"
        formatted += f"- **Feature**: {prd_info.get('feature', 'N/A')}\n"
        formatted += f"- **Test Point**: {prd_info.get('test_point', 'N/A')}\n"
        formatted += f"- **PRD Document**: {prd_info.get('prd_document', 'N/A')}\n"
        formatted += f"- **Steps**:\n```json\n{json.dumps(steps, ensure_ascii=False, indent=2)}\n```\n\n"
    
    return formatted


def get_jsonl_index_stats(jsonl_path: str) -> Dict[str, Any]:
    """Get statistics about JSONL data"""
    data = load_jsonl_data(jsonl_path)
    return {
        "exists": len(data) > 0,
        "num_examples": len(data),
        "path": jsonl_path
    }

